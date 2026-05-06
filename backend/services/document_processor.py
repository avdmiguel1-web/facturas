import asyncio
import json
import logging
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document as WordDocument
from dotenv import load_dotenv
from google import genai
from google.genai import types
from openpyxl import load_workbook
from PyPDF2 import PdfReader
import pdfplumber
# Nota: Para OCR en imágenes, se requiere instalar tesseract-ocr en el sistema.
# Descargar desde https://github.com/UB-Mannheim/tesseract/wiki
# Para Windows: choco install tesseract
# Para Linux: sudo apt install tesseract-ocr

from services.file_validation import FileValidationError, validate_saved_file

load_dotenv()

logger = logging.getLogger(__name__)


class DocumentProcessingError(RuntimeError):
    """Raised when the document cannot be processed consistently."""


@dataclass
class PreparedDocument:
    mode: str
    content: str
    mime_type: str
    file_path: Optional[str] = None


class DocumentProcessor:
    def __init__(self) -> None:
        self._disable_invalid_local_proxies()
        api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError("No se encontro GEMINI_API_KEY ni GOOGLE_API_KEY en el entorno.")

        self.client = genai.Client(api_key=api_key)
        primary_model = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash").strip()
        configured_fallbacks = [
            model.strip()
            for model in os.environ.get("GEMINI_FALLBACK_MODELS", "gemini-2.5-flash-lite,gemini-2.0-flash").split(",")
            if model.strip()
        ]
        self.models = [primary_model, *[model for model in configured_fallbacks if model != primary_model]]
        self.max_text_characters = int(os.environ.get("MAX_DOCUMENT_TEXT_CHARS", "120000"))
        self.max_rows_per_sheet = int(os.environ.get("MAX_SPREADSHEET_ROWS", "250"))
        self.retry_attempts = int(os.environ.get("GEMINI_RETRY_ATTEMPTS", "4"))

        self._classification_subtypes = {
            "Luz": ["luz", "electricidad", "ene"],
            "Agua": ["agua", "acueducto"],
            "Internet": ["internet", "wifi", "banda ancha", "isp", "datos"],
            "Mantenimiento": ["mantenimiento", "reparacion", "reparación", "servicio tecnico", "service"],
            "Telefonia": ["telefonia", "telefono", "telefónica", "movil", "móvil", "celular"],
            "Materiales": ["materiales", "insumos", "suministros", "herramientas", "papeleria", "papelería"],
            "Repuestos": ["repuestos", "refacciones", "refacción", "repuesto"],
            "Equipos": ["equipos", "equipo", "maquinaria", "hardware"],
            "Compras generales": ["compra", "compras", "factura de compra", "mercancia", "mercancía", "gasto"],
        }

        self._subcategory_to_category = {
            "Luz": "SERVICIOS",
            "Agua": "SERVICIOS",
            "Internet": "SERVICIOS",
            "Mantenimiento": "SERVICIOS",
            "Telefonia": "SERVICIOS",
            "Materiales": "CONSUMOS",
            "Repuestos": "CONSUMOS",
            "Equipos": "CONSUMOS",
            "Compras generales": "CONSUMOS",
        }

    def _disable_invalid_local_proxies(self) -> None:
        invalid_proxy_targets = {
            "http://127.0.0.1:9",
            "https://127.0.0.1:9",
            "http://localhost:9",
            "https://localhost:9",
            "127.0.0.1:9",
            "localhost:9",
        }

        proxy_keys = [
            "HTTP_PROXY",
            "HTTPS_PROXY",
            "ALL_PROXY",
            "http_proxy",
            "https_proxy",
            "all_proxy",
        ]

        for key in proxy_keys:
            value = (os.environ.get(key) or "").strip()
            if value and value in invalid_proxy_targets:
                logger.warning("Ignoring invalid proxy setting %s=%s for Gemini requests", key, value)
                os.environ.pop(key, None)

    async def process_document(self, file_path: str, module_type: str) -> Dict[str, Any]:
        metadata = validate_saved_file(file_path)

        if module_type == "specialized" and metadata["extension"] == ".pdf":
            parsed_pdf_data = await asyncio.to_thread(
                self._parse_specialized_pdf,
                Path(file_path),
            )
            if self._has_specialized_rows(parsed_pdf_data):
                normalized = self._normalize_extracted_data(parsed_pdf_data, module_type)
                return await self._classify_extracted_data(normalized, module_type)

        prepared = await asyncio.to_thread(
            self._prepare_document,
            Path(file_path),
            str(metadata["extension"]),
            str(metadata["file_type"]),
        )

        response_text = await self._generate_with_retries(prepared, module_type)
        parsed_data = self._parse_response(response_text)
        normalized = self._normalize_extracted_data(parsed_data, module_type)
        return await self._classify_extracted_data(normalized, module_type)

    async def _generate_with_retries(self, prepared: PreparedDocument, module_type: str) -> str:
        last_error: Optional[Exception] = None

        for model in self.models:
            for attempt in range(1, self.retry_attempts + 1):
                try:
                    return await asyncio.to_thread(
                        self._generate_content_sync,
                        prepared,
                        module_type,
                        model,
                    )
                except Exception as exc:  # pragma: no cover - network dependent
                    last_error = exc
                    logger.warning(
                        "Gemini request failed with model %s on attempt %s/%s: %s",
                        model,
                        attempt,
                        self.retry_attempts,
                        exc,
                    )

                    if self._is_model_unavailable_error(exc):
                        logger.warning("Skipping unsupported or unavailable model %s: %s", model, exc)
                        break

                    if not self._is_retryable_error(exc):
                        raise DocumentProcessingError(
                            f"No se pudo obtener una respuesta valida de Gemini: {exc}"
                        ) from exc

                    if attempt < self.retry_attempts:
                        await asyncio.sleep(min(2 ** attempt, 12))
                        continue

                    logger.warning("Switching to next Gemini model after exhausting retries for %s", model)
                    break

        if last_error and self._is_retryable_error(last_error):
            raise DocumentProcessingError(
                "Gemini esta temporalmente saturado. Se reintentaron varios modelos automaticamente, "
                "pero el servicio sigue respondiendo con alta demanda. Intenta nuevamente en unos minutos."
            ) from last_error

        raise DocumentProcessingError(
            f"No se pudo obtener una respuesta valida de Gemini: {last_error}"
        ) from last_error

    def _generate_content_sync(self, prepared: PreparedDocument, module_type: str, model: str) -> str:
        config = types.GenerateContentConfig(
            temperature=0,
            response_mime_type="application/json",
            response_schema=self._get_response_schema(module_type),
            system_instruction=self._get_system_instruction(module_type),
        )

        uploaded_file_name: Optional[str] = None

        try:
            if prepared.mode == "text":
                contents: List[Any] = [self._compose_text_prompt(module_type, prepared.content)]
            else:
                uploaded_file = self.client.files.upload(
                    file=prepared.file_path,
                    config=types.UploadFileConfig(
                        mime_type=prepared.mime_type,
                        display_name=Path(prepared.file_path or "").name,
                    ),
                )
                uploaded_file_name = uploaded_file.name
                contents = [uploaded_file, self._get_file_prompt(module_type)]

            response = self.client.models.generate_content(
                model=model,
                contents=contents,
                config=config,
            )
            return self._extract_response_text(response)
        finally:
            if uploaded_file_name:  # pragma: no cover - network dependent
                try:
                    self.client.files.delete(name=uploaded_file_name)
                except Exception as exc:
                    logger.warning("Could not delete uploaded Gemini file %s: %s", uploaded_file_name, exc)

    def _prepare_document(self, path: Path, extension: str, mime_type: str) -> PreparedDocument:
        if extension in {".png", ".jpg", ".jpeg"}:
            return PreparedDocument(mode="file", content="", mime_type=mime_type, file_path=str(path))

        if extension == ".pdf":
            extracted_text = self._extract_pdf_text(path)
            if self._has_usable_text(extracted_text):
                return PreparedDocument(
                    mode="text",
                    content=self._truncate_text("Documento convertido desde PDF.\n\n" + extracted_text),
                    mime_type="text/plain",
                )
            return PreparedDocument(mode="file", content="", mime_type=mime_type, file_path=str(path))

        if extension == ".docx":
            return PreparedDocument(
                mode="text",
                content=self._truncate_text(self._extract_docx_text(path)),
                mime_type="text/plain",
            )

        if extension in {".xlsx", ".xlsm"}:
            return PreparedDocument(
                mode="text",
                content=self._truncate_text(self._extract_spreadsheet_text(path)),
                mime_type="text/plain",
            )

        if extension in {".txt", ".csv"}:
            return PreparedDocument(
                mode="text",
                content=self._truncate_text(self._extract_text_file(path)),
                mime_type="text/plain",
            )

        raise FileValidationError(f"El archivo '{path.name}' no tiene un procesador configurado.")

    def _parse_specialized_pdf(self, path: Path) -> Dict[str, Any]:
        pages = self._extract_pdf_pages(path)
        sales: List[Dict[str, Any]] = []
        mobile_totals: List[Dict[str, Any]] = []
        rent_lines: List[Dict[str, Any]] = []
        consumption_lines: List[Dict[str, Any]] = []

        seen_sales: Set[Tuple[Optional[str], Optional[float], Optional[float]]] = set()
        seen_mobiles: Set[Tuple[Optional[str], Optional[str], Optional[float]]] = set()

        current_mobile: Optional[str] = None

        for page_text in pages:
            lines = self._get_page_lines(page_text)
            if not lines:
                continue

            mobile_number = self._extract_mobile_number(page_text)
            if mobile_number:
                current_mobile = mobile_number

            if current_mobile is None and self._contains_marker(page_text, "VENTAS DE TERCEROS") and not self._contains_marker(
                page_text, "VENTAS POR CUENTA DE TERCEROS"
            ):
                for item in self._parse_sales_summary_lines(lines):
                    key = (item["descripcion"], item["cantidad"], item["monto_bs"])
                    if key not in seen_sales:
                        seen_sales.add(key)
                        sales.append(item)

            if current_mobile and self._contains_marker(page_text, "RENTAS Y SERVICIOS"):
                for item in self._parse_rent_lines(lines):
                    rent_lines.append(item)

            if current_mobile and self._contains_marker(page_text, "RESUMEN DE CONSUMO"):
                for item in self._parse_consumption_lines(lines):
                    consumption_lines.append(item)

            total_facturado = self._extract_total_facturado_movil(page_text)
            if current_mobile and total_facturado is not None:
                item = {
                    "numero_movil": current_mobile,
                    "descripcion": f"No Movil {current_mobile}",
                    "monto_bs": total_facturado,
                }
                key = (item["numero_movil"], item["descripcion"], item["monto_bs"])
                if key not in seen_mobiles:
                    seen_mobiles.add(key)
                    mobile_totals.append(item)

        return {
            "ventas_terceros": sales,
            "telefonia_movil": mobile_totals,
            "rentas_servicios": rent_lines,
            "resumen_consumo": consumption_lines,
        }

    def _extract_pdf_pages(self, path: Path) -> List[str]:
        try:
            reader = PdfReader(str(path))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    return []

            pages: List[str] = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text)
            return pages
        except Exception as exc:
            logger.warning("Could not extract PDF pages from %s: %s", path, exc)
            return []

    def _extract_pdf_text(self, path: Path) -> str:
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for index, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(f"Pagina {index}\n{page_text.strip()}")
                return "\n\n".join(pages)
        except Exception as exc:
            logger.warning("Could not extract PDF text from %s: %s", path, exc)
            return ""

    def _get_page_lines(self, page_text: str) -> List[str]:
        return [line.strip() for line in page_text.splitlines() if line.strip()]

    def _contains_marker(self, text: str, marker: str) -> bool:
        normalized_text = self._normalize_marker_text(text)
        normalized_marker = self._normalize_marker_text(marker)
        return normalized_marker in normalized_text

    def _normalize_marker_text(self, text: str) -> str:
        replacements = str.maketrans(
            {
                "Á": "A",
                "É": "E",
                "Í": "I",
                "Ó": "O",
                "Ú": "U",
                "á": "a",
                "é": "e",
                "í": "i",
                "ó": "o",
                "ú": "u",
            }
        )
        return text.translate(replacements).upper()

    def _extract_mobile_number(self, page_text: str) -> Optional[str]:
        match = re.search(r"No\s+M[oó]vil\s+(\d{7,})", page_text, flags=re.IGNORECASE)
        if match:
            return match.group(1)

        match = re.search(r"M[oó]vil:\s*(\d{7,})", page_text, flags=re.IGNORECASE)
        if match:
            return match.group(1)

        return None

    def _extract_total_facturado_movil(self, page_text: str) -> Optional[float]:
        match = re.search(
            r"Total Facturado\s+del mes\s+TELEFON[IÍ]A\s+M[ÓO]VIL\s+([0-9.,']+)",
            page_text,
            flags=re.IGNORECASE,
        )
        if not match:
            return None
        return self._coerce_number(match.group(1))

    def _parse_sales_summary_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        in_section = False

        for line in lines:
            if self._contains_marker(line, "VENTAS POR CUENTA DE TERCEROS"):
                continue

            if self._contains_marker(line, "VENTAS DE TERCEROS"):
                in_section = True
                continue

            if not in_section:
                continue

            normalized = self._normalize_marker_text(line)
            if (
                "TELEFONIA MOVIL" in normalized
                or "RESUMEN DE CONSUMO" in normalized
                or normalized.startswith("SE EMITE")
            ):
                break

            if (
                "DESCRIPCION" in normalized
                or normalized.startswith("SUB-TOTAL")
                or normalized.startswith("I.C.S.")
                or normalized.startswith("I.V.A.")
                or normalized.startswith("TOTAL ")
            ):
                continue

            match = re.match(r"^(?P<descripcion>.+?)\s+(?P<cantidad>\d+(?:[.,]\d+)?)\s+(?P<monto>[0-9.,']+)$", line)
            if not match:
                continue

            rows.append(
                {
                    "descripcion": self._clean_string(self._strip_account_code(match.group("descripcion"))),
                    "cantidad": self._coerce_number(match.group("cantidad")),
                    "monto_bs": self._coerce_number(match.group("monto")),
                }
            )

        return rows

    def _parse_rent_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        in_section = False

        for line in lines:
            if self._contains_marker(line, "RENTAS Y SERVICIOS"):
                in_section = True
                continue

            if not in_section:
                continue

            normalized = self._normalize_marker_text(line)
            if "RESUMEN DE CONSUMO" in normalized:
                break

            if (
                "DESCRIPCION" in normalized
                or "DESDEFECHA" in normalized
                or "HASTAMONTO" in normalized
                or normalized.startswith("SUB-TOTAL")
                or normalized.startswith("IVA ")
                or normalized.startswith("IVA(")
                or normalized.startswith("TOTAL FACTURADO")
            ):
                continue

            match = re.match(
                r"^(?P<descripcion>.+?)\s+(?P<fecha_desde>\d{2}/\d{2}/\d{4})\s+(?P<fecha_hasta>\d{2}/\d{2}/\d{4})\s+(?P<monto>[0-9.,']+)$",
                line,
            )
            if not match:
                continue

            rows.append(
                {
                    "descripcion": self._clean_string(self._strip_account_code(match.group("descripcion"))),
                    "fecha_desde": match.group("fecha_desde"),
                    "fecha_hasta": match.group("fecha_hasta"),
                    "monto_bs": self._coerce_number(match.group("monto")),
                }
            )

        return rows

    def _parse_consumption_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        in_section = False

        for line in lines:
            if self._contains_marker(line, "RESUMEN DE CONSUMO"):
                in_section = True
                continue

            if not in_section:
                continue

            normalized = self._normalize_marker_text(line)
            if (
                normalized.startswith("SUB-TOTAL")
                or normalized.startswith("IVA ")
                or normalized.startswith("IVA(")
                or normalized.startswith("TOTAL FACTURADO")
            ):
                continue

            if "DESCRIPCION" in normalized or "CONSUMIDASMONTO" in normalized:
                continue

            match = re.match(
                r"^(?P<descripcion>.+?)\s+(?P<unidades>[0-9.,']+)\s+(?P<monto>[0-9.,']+)$",
                line,
            )
            if not match:
                continue

            rows.append(
                {
                    "descripcion": self._clean_string(self._strip_account_code(match.group("descripcion"))),
                    "unidades_consumidas": self._coerce_number(match.group("unidades").replace("'", ".")),
                    "monto_bs": self._coerce_number(match.group("monto")),
                }
            )

        return rows

    def _strip_account_code(self, description: str) -> str:
        return re.sub(r"^\d{3,5}\s+", "", description).strip()

    def _has_specialized_rows(self, data: Dict[str, Any]) -> bool:
        return any(
            self._ensure_list(data.get(section))
            for section in ("ventas_terceros", "telefonia_movil", "rentas_servicios", "resumen_consumo")
        )

    def _extract_docx_text(self, path: Path) -> str:
        document = WordDocument(str(path))
        chunks: List[str] = []

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        if paragraphs:
            chunks.append("Parrafos\n" + "\n".join(paragraphs))

        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                chunks.append(f"Tabla {table_index}\n" + "\n".join(rows))

        if not chunks:
            raise DocumentProcessingError(f"El archivo Word '{path.name}' no contiene texto utilizable.")

        return "\n\n".join(chunks)

    def _extract_spreadsheet_text(self, path: Path) -> str:
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        chunks: List[str] = []

        try:
            for sheet in workbook.worksheets:
                rows: List[str] = []
                non_empty_rows = 0

                for row in sheet.iter_rows(values_only=True):
                    values = [self._stringify_cell(value) for value in row]
                    if not any(values):
                        continue

                    non_empty_rows += 1
                    rows.append("\t".join(values))
                    if non_empty_rows >= self.max_rows_per_sheet:
                        rows.append("[Filas adicionales omitidas para mantener la peticion eficiente]")
                        break

                if rows:
                    chunks.append(f"Hoja: {sheet.title}\n" + "\n".join(rows))
        finally:
            workbook.close()

        if not chunks:
            raise DocumentProcessingError(f"El archivo Excel '{path.name}' no contiene celdas con datos.")

        return "\n\n".join(chunks)

    def _extract_text_file(self, path: Path) -> str:
        file_bytes = path.read_bytes()
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                text = file_bytes.decode(encoding)
                if text.strip():
                    return text
            except UnicodeDecodeError:
                continue

        raise DocumentProcessingError(f"El archivo '{path.name}' no pudo leerse como texto.")

    def _compose_text_prompt(self, module_type: str, extracted_text: str) -> str:
        return (
            f"{self._get_user_prompt(module_type)}\n\n"
            "Documento normalizado para analisis:\n"
            f"{extracted_text}"
        )

    def _get_file_prompt(self, module_type: str) -> str:
        return self._get_user_prompt(module_type) + "\n\nAnaliza visualmente el archivo adjunto y responde solo con JSON valido."

    def _get_system_instruction(self, module_type: str) -> str:
        if module_type == "general":
            return (
                "Eres un extractor de datos de facturas. Devuelves JSON limpio, consistente "
                "y sin texto adicional. Nunca inventas valores ausentes."
            )

        if module_type == "specialized":
            return (
                "Eres un extractor de datos tabulares para analisis de gastos. Devuelves "
                "JSON limpio, consistente y sin texto adicional. Nunca inventas filas."
            )

        raise ValueError(f"Tipo de modulo desconocido: {module_type}")

    def _get_user_prompt(self, module_type: str) -> str:
        if module_type == "general":
            return (
                "Extrae los campos de facturacion con estas reglas:\n"
                "- Responde solo con un objeto JSON.\n"
                "- Usa null cuando el valor no exista en el documento.\n"
                "- Conserva fechas y textos como aparecen en el documento.\n"
                "- Convierte importes numericos a numero decimal sin simbolos de moneda.\n"
                "- Campos esperados: proveedor, rif, fecha, periodo, subtotal, iva, monto_total, moneda, tipo."
            )

        if module_type == "specialized":
            return (
                "Extrae las secciones del documento con estas reglas:\n"
                "- Responde solo con un objeto JSON.\n"
                "- Usa arrays vacios cuando una seccion no exista.\n"
                "- No agregues filas inventadas.\n"
                "- Convierte montos y cantidades a numero decimal.\n"
                "- Secciones esperadas: ventas_terceros, telefonia_movil, rentas_servicios, resumen_consumo."
            )

        raise ValueError(f"Tipo de modulo desconocido: {module_type}")

    def _get_response_schema(self, module_type: str) -> Dict[str, Any]:
        if module_type == "general":
            return {
                "type": "object",
                "properties": {
                    "proveedor": {"type": "string"},
                    "rif": {"type": "string"},
                    "fecha": {"type": "string"},
                    "periodo": {"type": "string"},
                    "subtotal": {"type": "number"},
                    "iva": {"type": "number"},
                    "monto_total": {"type": "number"},
                    "moneda": {"type": "string"},
                    "tipo": {"type": "string"},
                },
            }

        if module_type == "specialized":
            line_schema = {
                "type": "object",
                "properties": {
                    "descripcion": {"type": "string"},
                    "monto_bs": {"type": "number"},
                },
            }
            return {
                "type": "object",
                "properties": {
                    "ventas_terceros": {
                        "type": "array",
                        "items": {
                            **line_schema,
                            "properties": {
                                "descripcion": {"type": "string"},
                                "cantidad": {"type": "number"},
                                "monto_bs": {"type": "number"},
                            },
                        },
                    },
                    "telefonia_movil": {
                        "type": "array",
                        "items": {
                            **line_schema,
                            "properties": {
                                "numero_movil": {"type": "string"},
                                "descripcion": {"type": "string"},
                                "monto_bs": {"type": "number"},
                            },
                        },
                    },
                    "rentas_servicios": {
                        "type": "array",
                        "items": {
                            **line_schema,
                            "properties": {
                                "descripcion": {"type": "string"},
                                "fecha_desde": {"type": "string"},
                                "fecha_hasta": {"type": "string"},
                                "monto_bs": {"type": "number"},
                            },
                        },
                    },
                    "resumen_consumo": {
                        "type": "array",
                        "items": {
                            **line_schema,
                            "properties": {
                                "descripcion": {"type": "string"},
                                "unidades_consumidas": {"type": "number"},
                                "monto_bs": {"type": "number"},
                            },
                        },
                    },
                },
            }

        raise ValueError(f"Tipo de modulo desconocido: {module_type}")

    def _extract_response_text(self, response: Any) -> str:
        parsed = getattr(response, "parsed", None)
        if parsed is not None:
            return json.dumps(parsed, ensure_ascii=False)

        text = getattr(response, "text", None)
        if text:
            return text

        raise DocumentProcessingError("Gemini no devolvio texto ni JSON parseable.")

    def _parse_response(self, response_text: str) -> Dict[str, Any]:
        cleaned = response_text.strip()

        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.DOTALL).strip()

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError as exc:
            logger.error("Gemini response was not valid JSON: %s", cleaned)
            raise DocumentProcessingError("La respuesta de Gemini no fue un JSON valido.") from exc

        if not isinstance(data, dict):
            raise DocumentProcessingError("La respuesta de Gemini no contiene un objeto JSON.")

        return data

    def _normalize_extracted_data(self, data: Dict[str, Any], module_type: str) -> Dict[str, Any]:
        if module_type == "general":
            return {
                "proveedor": self._clean_string(data.get("proveedor")),
                "rif": self._clean_string(data.get("rif")),
                "fecha": self._clean_string(data.get("fecha")),
                "periodo": self._clean_string(data.get("periodo")),
                "subtotal": self._coerce_number(data.get("subtotal")),
                "iva": self._coerce_number(data.get("iva")),
                "monto_total": self._coerce_number(data.get("monto_total")),
                "moneda": self._clean_string(data.get("moneda")),
                "tipo": self._clean_string(data.get("tipo")),
            }

        if module_type == "specialized":
            return {
                "ventas_terceros": [
                    {
                        "descripcion": self._clean_string(item.get("descripcion")),
                        "cantidad": self._coerce_number(item.get("cantidad")),
                        "monto_bs": self._coerce_number(item.get("monto_bs")),
                    }
                    for item in self._ensure_list(data.get("ventas_terceros"))
                ],
                "telefonia_movil": [
                    {
                        "numero_movil": self._clean_string(item.get("numero_movil")),
                        "descripcion": self._clean_string(item.get("descripcion")),
                        "monto_bs": self._coerce_number(item.get("monto_bs")),
                    }
                    for item in self._ensure_list(data.get("telefonia_movil"))
                ],
                "rentas_servicios": [
                    {
                        "descripcion": self._clean_string(item.get("descripcion")),
                        "fecha_desde": self._clean_string(item.get("fecha_desde")),
                        "fecha_hasta": self._clean_string(item.get("fecha_hasta")),
                        "monto_bs": self._coerce_number(item.get("monto_bs")),
                    }
                    for item in self._ensure_list(data.get("rentas_servicios"))
                ],
                "resumen_consumo": [
                    {
                        "descripcion": self._clean_string(item.get("descripcion")),
                        "unidades_consumidas": self._coerce_number(item.get("unidades_consumidas")),
                        "monto_bs": self._coerce_number(item.get("monto_bs")),
                    }
                    for item in self._ensure_list(data.get("resumen_consumo"))
                ],
            }

        raise ValueError(f"Tipo de modulo desconocido: {module_type}")

    async def _classify_extracted_data(self, data: Dict[str, Any], module_type: str) -> Dict[str, Any]:
        if module_type == "general":
            text_to_classify = f"Proveedor: {data.get('proveedor', '')} Tipo: {data.get('tipo', '')}"
            categoria_tipo, categoria_subtipo = await self._classify_text(text_to_classify)
            data["categoria_tipo"] = categoria_tipo
            data["categoria_subtipo"] = categoria_subtipo
            return data

        if module_type == "specialized":
            for category in ["ventas_terceros", "telefonia_movil", "rentas_servicios", "resumen_consumo"]:
                if category in data:
                    for item in data[category]:
                        text_to_classify = item.get("descripcion", "")
                        categoria_tipo, categoria_subtipo = await self._classify_text(text_to_classify)
                        item["categoria_tipo"] = categoria_tipo
                        item["categoria_subtipo"] = categoria_subtipo
            return data

        return data

    def _normalize_classification_text(self, text: str) -> str:
        normalized = re.sub(r"[^a-zA-Z0-9\s]", " ", text or "")
        return re.sub(r"\s+", " ", normalized).strip().lower()

    def _classify_with_keywords(self, text: str) -> Tuple[Optional[str], Optional[str]]:
        normalized = self._normalize_classification_text(text)
        for subtype, keywords in self._classification_subtypes.items():
            for keyword in keywords:
                if keyword.lower() in normalized:
                    return self._subcategory_to_category.get(subtype, "DESCONOCIDO"), subtype

        for subtype, category in self._subcategory_to_category.items():
            if subtype.lower() in normalized:
                return category, subtype

        return None, None

    async def _classify_text(self, text: str) -> Tuple[str, str]:
        category, subtype = self._classify_with_keywords(text)
        if category:
            return category, subtype or "GENERAL"

        prompt = f"""
Clasifica el siguiente texto y responde en un objeto JSON.

Reglas:
- categoria_tipo debe ser SERVICIOS o CONSUMOS.
- categoria_subtipo debe ser una subcategoria como Luz, Agua, Internet, Mantenimiento, Repuestos, Equipos, Materiales, Compras generales, Telefonia.
- Si no puedes identificar un subtipo preciso, usa "GENERAL".

Texto: {text}

Responde solo con JSON valido. Ejemplo:
{{"categoria_tipo": "SERVICIOS", "categoria_subtipo": "Internet"}}
"""
        try:
            response = await asyncio.to_thread(
                self.client.models.generate_content,
                model=self.models[0],
                contents=[prompt],
                config=types.GenerateContentConfig(temperature=0),
            )
            result_text = self._extract_response_text(response).strip()
            if result_text.startswith("```"):
                result_text = re.sub(r"^```(?:json)?\s*|\s*```$", "", result_text, flags=re.DOTALL).strip()

            parsed = json.loads(result_text)
            category = str(parsed.get("categoria_tipo", "DESCONOCIDO")).strip().upper()
            subtype = str(parsed.get("categoria_subtipo", "GENERAL")).strip() or "GENERAL"
            if category not in {"SERVICIOS", "CONSUMOS"}:
                category = "DESCONOCIDO"
            return category, subtype
        except Exception as e:
            logger.warning(f"Error clasificando texto: {e}")
            return "DESCONOCIDO", "GENERAL"

    def _truncate_text(self, text: str) -> str:
        normalized = text.strip()
        if not normalized:
            raise DocumentProcessingError("El documento no contiene texto utilizable.")

        if len(normalized) <= self.max_text_characters:
            return normalized

        return normalized[: self.max_text_characters] + "\n\n[Contenido truncado para mantener la solicitud dentro del limite operativo]"

    def _has_usable_text(self, text: str) -> bool:
        return len(re.sub(r"\s+", "", text or "")) >= 80

    def _stringify_cell(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value).strip()

    def _ensure_list(self, value: Any) -> List[Dict[str, Any]]:
        if not value or not isinstance(value, list):
            return []
        return [item for item in value if isinstance(item, dict)]

    def _clean_string(self, value: Any) -> Optional[str]:
        if value is None:
            return None
        text = str(value).strip()
        return text or None

    def _coerce_number(self, value: Any) -> Optional[float]:
        if value is None or value == "":
            return None
        if isinstance(value, (int, float)):
            return float(value)

        text = str(value).strip()
        if not text:
            return None

        text = re.sub(r"[^0-9,.\-]", "", text)
        if text.count(",") > 0 and text.count(".") > 0:
            if text.rfind(",") > text.rfind("."):
                text = text.replace(".", "").replace(",", ".")
            else:
                text = text.replace(",", "")
        elif text.count(",") > 0:
            text = text.replace(".", "").replace(",", ".")

        try:
            return float(text)
        except ValueError:
            return None

    def _is_retryable_error(self, error: Exception) -> bool:
        message = str(error).upper()
        retry_markers = [
            "503",
            "504",
            "429",
            "UNAVAILABLE",
            "RESOURCE_EXHAUSTED",
            "DEADLINE_EXCEEDED",
            "HIGH DEMAND",
        ]
        return any(marker in message for marker in retry_markers)

    def _is_model_unavailable_error(self, error: Exception) -> bool:
        message = str(error).upper()
        return "404" in message and "NOT_FOUND" in message and "MODEL" in message
